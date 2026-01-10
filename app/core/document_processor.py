"""
KMU Knowledge Assistant - Dokumentenverarbeitung
Unterstützt PDF, DOCX, TXT, MD, RTF, XLSX, CSV, MSG, EML, Bilder
"""

import io
import re
import csv
import email
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field
from datetime import datetime
import hashlib

# Document processing libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import extract_msg
    EXTRACT_MSG_AVAILABLE = True
except ImportError:
    EXTRACT_MSG_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from app.config import config, SUPPORTED_FORMATS


@dataclass
class DocumentChunk:
    """Ein Chunk eines Dokuments"""
    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None


@dataclass
class ProcessedDocument:
    """Verarbeitetes Dokument"""
    id: str
    filename: str
    file_type: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    raw_text: str
    processed_at: datetime = field(default_factory=datetime.now)
    
    @property
    def chunk_count(self) -> int:
        return len(self.chunks)
    
    @property
    def character_count(self) -> int:
        return len(self.raw_text)


class DocumentProcessor:
    """Verarbeitet verschiedene Dokumentformate"""
    
    def __init__(self):
        self.chunk_size = config.rag.chunk_size
        self.chunk_overlap = config.rag.chunk_overlap
    
    def process_file(
        self,
        file_path: Path | str,
        knowledge_base_id: str,
        uploader_id: str = "system"
    ) -> ProcessedDocument:
        """Verarbeitet eine Datei und erstellt Chunks"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Datei nicht gefunden: {file_path}")
        
        extension = file_path.suffix.lower()
        
        # Text extrahieren
        raw_text = self._extract_text(file_path, extension)
        
        # Dokument-ID generieren
        doc_id = self._generate_doc_id(file_path, raw_text)
        
        # Content-Hash für Change-Detection (SHA256)
        content_hash = hashlib.sha256(raw_text.encode()).hexdigest()

        # Metadata
        metadata = {
            "filename": file_path.name,
            "file_type": extension,
            "file_size": file_path.stat().st_size,
            "knowledge_base": knowledge_base_id,
            "uploader": uploader_id,
            "upload_date": datetime.now().isoformat(),
            "content_hash": content_hash
        }
        
        # Chunks erstellen
        chunks = self._create_chunks(raw_text, doc_id, metadata)
        
        return ProcessedDocument(
            id=doc_id,
            filename=file_path.name,
            file_type=extension,
            chunks=chunks,
            metadata=metadata,
            raw_text=raw_text
        )
    
    def process_bytes(
        self,
        content: bytes,
        filename: str,
        knowledge_base_id: str,
        uploader_id: str = "system"
    ) -> ProcessedDocument:
        """Verarbeitet Bytes direkt (für Upload)"""
        extension = Path(filename).suffix.lower()
        
        # Text extrahieren
        raw_text = self._extract_text_from_bytes(content, extension)
        
        # Dokument-ID generieren
        doc_id = self._generate_doc_id_from_bytes(content, filename)
        
        # Content-Hash für Change-Detection (SHA256)
        content_hash = hashlib.sha256(raw_text.encode()).hexdigest()

        # Metadata
        metadata = {
            "filename": filename,
            "file_type": extension,
            "file_size": len(content),
            "knowledge_base": knowledge_base_id,
            "uploader": uploader_id,
            "upload_date": datetime.now().isoformat(),
            "content_hash": content_hash
        }
        
        # Chunks erstellen
        chunks = self._create_chunks(raw_text, doc_id, metadata)
        
        return ProcessedDocument(
            id=doc_id,
            filename=filename,
            file_type=extension,
            chunks=chunks,
            metadata=metadata,
            raw_text=raw_text
        )
    
    def _extract_text(self, file_path: Path, extension: str) -> str:
        """Extrahiert Text aus einer Datei"""
        with open(file_path, "rb") as f:
            content = f.read()
        return self._extract_text_from_bytes(content, extension)
    
    def _extract_text_from_bytes(self, content: bytes, extension: str) -> str:
        """Extrahiert Text aus Bytes"""
        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt": self._extract_txt,
            ".md": self._extract_txt,
            ".rtf": self._extract_rtf,
            ".xlsx": self._extract_xlsx,
            ".csv": self._extract_csv,
            ".msg": self._extract_msg,
            ".eml": self._extract_eml,
            ".html": self._extract_html,
            ".htm": self._extract_html,
            ".png": self._extract_image,
            ".jpg": self._extract_image,
            ".jpeg": self._extract_image,
            ".tiff": self._extract_image
        }
        
        extractor = extractors.get(extension)
        if extractor:
            return extractor(content)
        else:
            raise ValueError(f"Nicht unterstütztes Format: {extension}")
    
    def _extract_pdf(self, content: bytes) -> str:
        """Extrahiert Text aus PDF (mit OCR-Fallback für Bild-PDFs)"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF nicht installiert. Bitte 'pip install pymupdf' ausführen.")

        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[Seite {page_num}]\n{text}")
            elif OCR_AVAILABLE:
                # Fallback: OCR für Bild-PDFs
                try:
                    # Seite als Bild rendern (300 DPI)
                    mat = fitz.Matrix(300/72, 300/72)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")

                    # OCR durchführen
                    image = Image.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image, lang="deu+eng")
                    if ocr_text.strip():
                        text_parts.append(f"[Seite {page_num} (OCR)]\n{ocr_text}")
                except Exception as e:
                    print(f"OCR-Fehler auf Seite {page_num}: {e}")

        doc.close()
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, content: bytes) -> str:
        """Extrahiert Text aus DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx nicht installiert. Bitte 'pip install python-docx' ausführen.")
        
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
    
    def _extract_txt(self, content: bytes) -> str:
        """Extrahiert Text aus TXT/MD"""
        # Versuche verschiedene Encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")
    
    def _extract_rtf(self, content: bytes) -> str:
        """Extrahiert Text aus RTF (Basic)"""
        text = content.decode("utf-8", errors="replace")
        # Einfache RTF-Tag-Entfernung
        text = re.sub(r'\\[a-z]+\d*\s?', '', text)
        text = re.sub(r'[{}]', '', text)
        return text.strip()
    
    def _extract_xlsx(self, content: bytes) -> str:
        """Extrahiert Text aus XLSX"""
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl nicht installiert. Bitte 'pip install openpyxl' ausführen.")
        
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"[Tabellenblatt: {sheet_name}]")
            
            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    rows.append(row_text)
            
            text_parts.append("\n".join(rows))
        
        wb.close()
        return "\n\n".join(text_parts)
    
    def _extract_csv(self, content: bytes) -> str:
        """Extrahiert Text aus CSV"""
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = [" | ".join(row) for row in reader]
        return "\n".join(rows)
    
    def _extract_msg(self, content: bytes) -> str:
        """Extrahiert Text aus MSG (Outlook)"""
        if not EXTRACT_MSG_AVAILABLE:
            raise ImportError("extract-msg nicht installiert. Bitte 'pip install extract-msg' ausführen.")
        
        msg = extract_msg.Message(io.BytesIO(content))
        parts = [
            f"Von: {msg.sender}",
            f"An: {msg.to}",
            f"Betreff: {msg.subject}",
            f"Datum: {msg.date}",
            "",
            msg.body or ""
        ]
        return "\n".join(parts)
    
    def _extract_eml(self, content: bytes) -> str:
        """Extrahiert Text aus EML"""
        msg = email.message_from_bytes(content)
        
        parts = [
            f"Von: {msg.get('From', '')}",
            f"An: {msg.get('To', '')}",
            f"Betreff: {msg.get('Subject', '')}",
            f"Datum: {msg.get('Date', '')}",
            ""
        ]
        
        # Body extrahieren
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                parts.append(payload.decode("utf-8", errors="replace"))
        
        return "\n".join(parts)
    
    def _extract_html(self, content: bytes) -> str:
        """Extrahiert Text aus HTML"""
        if not BS4_AVAILABLE:
            raise ImportError("BeautifulSoup nicht installiert. Bitte 'pip install beautifulsoup4' ausführen.")
        
        soup = BeautifulSoup(content, "html.parser")
        
        # Script und Style entfernen
        for script in soup(["script", "style"]):
            script.decompose()
        
        return soup.get_text(separator="\n", strip=True)
    
    def _extract_image(self, content: bytes) -> str:
        """Extrahiert Text aus Bildern via OCR"""
        if not OCR_AVAILABLE:
            raise ImportError(
                "OCR nicht verfügbar. Bitte 'pip install pytesseract pillow' "
                "und Tesseract installieren."
            )
        
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image, lang="deu+eng")
        return text
    
    def _create_chunks(
        self,
        text: str,
        doc_id: str,
        base_metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Erstellt überlappende Chunks aus Text"""
        if not text.strip():
            return []
        
        # Normalisiere Whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Schätze Token-Länge (ca. 4 Zeichen pro Token für Deutsch)
        chars_per_token = 4
        chunk_chars = self.chunk_size * chars_per_token
        overlap_chars = self.chunk_overlap * chars_per_token
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_chars
            
            # Versuche an Satzende zu brechen
            if end < len(text):
                # Suche nach Satzende in den letzten 20% des Chunks
                search_start = start + int(chunk_chars * 0.8)
                search_region = text[search_start:end]
                
                # Finde letztes Satzende
                for sep in [". ", "! ", "? ", "\n"]:
                    last_sep = search_region.rfind(sep)
                    if last_sep != -1:
                        end = search_start + last_sep + len(sep)
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_id = f"{doc_id}_chunk_{chunk_index}"
                
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": chunk_index,
                    "chunk_start": start,
                    "chunk_end": end
                })
                
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk_text,
                    metadata=chunk_metadata
                ))
                
                chunk_index += 1
            
            # Nächster Start mit Überlappung
            start = end - overlap_chars
            if start <= 0 and chunk_index > 0:
                start = end
        
        return chunks
    
    def _generate_doc_id(self, file_path: Path, content: str) -> str:
        """Generiert eindeutige Dokument-ID"""
        hash_input = f"{file_path.name}_{len(content)}_{content[:1000]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:16]
    
    def _generate_doc_id_from_bytes(self, content: bytes, filename: str) -> str:
        """Generiert eindeutige Dokument-ID aus Bytes"""
        hash_input = f"{filename}_{len(content)}_{content[:1000]}"
        return hashlib.md5(hash_input.encode() if isinstance(hash_input, str) else hash_input).hexdigest()[:16]


# Globale Instanz
document_processor = DocumentProcessor()
